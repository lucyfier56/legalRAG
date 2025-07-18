
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from groq import Groq
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz
from typing import List, Dict, Any
import logging
import base64
from io import BytesIO
import os
import tempfile
import pytesseract
from PIL import Image
import glob
from pathlib import Path
import re
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

GROQ_API_KEY = ""

class RateLimiter:
    def __init__(self, max_requests_per_minute=30):
        self.max_requests = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            # Remove requests older than 1 minute
            self.requests = [req_time for req_time in self.requests if now - req_time < 60]
            
            if len(self.requests) >= self.max_requests:
                # Wait until the oldest request is more than 1 minute old
                sleep_time = 60 - (now - self.requests[0]) + 0.1
                if sleep_time > 0:
                    time.sleep(sleep_time)
                    # Clean up again after sleeping
                    now = time.time()
                    self.requests = [req_time for req_time in self.requests if now - req_time < 60]
            
            self.requests.append(now)

class TimelineExtractor:
    def __init__(self):
        self.client = Groq(api_key=GROQ_API_KEY)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.logger = logging.getLogger(__name__)
        self.rate_limiter = RateLimiter(max_requests_per_minute=30)

    def clean_date_string(self, date_str: str) -> str:
        """Clean date string by removing unwanted characters and symbols"""
        # Remove any leading/trailing whitespace
        cleaned = date_str.strip()
        
        # Remove markdown formatting, asterisks, dashes, and other symbols
        cleaned = re.sub(r'[*\-_#`~]', '', cleaned)
        
        # Remove any leading/trailing non-alphanumeric characters except date separators
        cleaned = re.sub(r'^[^\w\d/,.\s]+|[^\w\d/,.\s]+$', '', cleaned)
        
        # Remove any remaining unwanted characters but keep date separators
        cleaned = re.sub(r'[^\w\d\s/,.-]', '', cleaned)
        
        # Remove multiple spaces and normalize
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # Remove any remaining leading/trailing spaces
        cleaned = cleaned.strip()
        
        return cleaned

    def filter_main_content(self, text: str) -> str:
        """Filter out page headers, footers, and metadata to focus on main content"""
        lines = text.split('\n')
        filtered_lines = []
        
        # Keywords that indicate metadata/headers/footers to exclude
        exclude_keywords = [
            'digitally signed', 'digital signing', 'signing date', 'signed by',
            'page number', 'page', 'footer', 'header',
            'timestamp', 'generated on', 'printed on',
            'file name', 'document id', 'version',
            'watermark', 'confidential', 'draft'
        ]
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Skip empty lines
            if not line_lower:
                continue
                
            # Skip lines that are likely metadata/headers/footers
            if any(keyword in line_lower for keyword in exclude_keywords):
                continue
                
            # Skip lines that are just dates without context (likely metadata)
            if re.match(r'^\s*\d{1,2}[./\-]\d{1,2}[./\-]\d{4}\s*$', line.strip()):
                continue
                
            # Skip very short lines (likely page numbers or metadata)
            if len(line.strip()) < 10:
                continue
                
            filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)

    def extract_dates_and_events_batch(self, page_data_list: List[Dict]) -> List[Dict]:
        """Extract dates and events from multiple pages in a single API call"""
        if not page_data_list:
            return []
        
        # Apply rate limiting
        self.rate_limiter.wait_if_needed()
        
        self.logger.info(f"Processing batch of {len(page_data_list)} pages")
        
        # Prepare batch content
        batch_content = []
        for page_data in page_data_list:
            filtered_text = self.filter_main_content(page_data['text'])
            batch_content.append({
                'page_number': page_data['page_number'],
                'content': filtered_text
            })
        
        # Create batch prompt
        batch_prompt = f"""
        You are an expert legal document analyst. Analyze the following MAIN CONTENT from multiple pages and extract ONLY substantive case-related dates and events.

        CRITICAL FILTERING RULES:
        1. IGNORE all document metadata, digital signatures, timestamps, and page information
        2. IGNORE dates related to "digitally signed", "signing date", "generated on", "printed on"
        3. FOCUS ONLY on case-related events like hearings, filings, orders, judgments, applications
        4. ONLY extract dates that have substantive legal significance to the case

        CONTENT FOCUS AREAS:
        - Court proceedings and hearings
        - Filing of petitions, applications, appeals
        - Court orders and judgments  
        - Legal deadlines and compliance dates
        - Case transfers and assignments
        - Settlement negotiations and outcomes

        DATE FORMATS TO RECOGNIZE:
        - Complete dates: "14th February 2017" → "14/02/2017"
        - Year only: "1996" → "1996" (NOT "01/01/1996")
        - Case numbers with years: "W.P.(C) 915/1996" → "1996"

        SPECIAL ATTENTION TO MULTIPLE DATES:
        - When multiple dates appear in parentheses like "(20th December, 2015, 18th February, 2016, 09th July, 2017 as well as 02nd March, 2018)", extract EACH date separately
        - Apply the SAME event description to ALL dates mentioned in the same context
        - Look for patterns like "incidents dated X, Y, and Z" or "on dates (A, B, C)"

        EVENT DESCRIPTION REQUIREMENTS:
        - Provide detailed, meaningful descriptions of what actually happened
        - Include specific legal actions, court decisions, or procedural steps
        - Mention parties involved (petitioner, respondent, court, etc.)
        - Include case numbers, section references, and legal citations when available
        - Describe the nature and outcome of proceedings

        FORMAT FOR EACH PAGE: 
        PAGE [page_number]:
        DATE | Detailed event description with legal significance

        EXAMPLES OF GOOD EXTRACTIONS:
        - "15/03/2020 | High Court issued interim stay order in W.P.(C) 1234/2020 restraining respondent from taking coercive action against petitioner"
        - "1996 | Filing of original Writ Petition (C) No. 915/1996 challenging termination of service"
        - "22/08/2019 | Court directed petitioner to file additional affidavit within two weeks regarding compliance with environmental norms"

        EXAMPLES TO AVOID:
        - Digital signing information
        - Document generation dates
        - Page numbers or metadata
        - Generic "event occurred" descriptions

        Multiple pages content to analyze:
        """
        
        for page_data in batch_content:
            batch_prompt += f"\n\nPAGE {page_data['page_number']}:\n{page_data['content']}"
        
        batch_prompt += "\n\nExtract ONLY substantive case-related events with detailed descriptions for each page."
        
        try:
            self.logger.info("Sending batch request to Groq API")
            response = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": """You are a specialized legal document analyst focused on extracting substantive case information. 
                        
                        Your expertise:
                        - Distinguishing between case content and document metadata
                        - Identifying legally significant events and proceedings
                        - Providing detailed, meaningful event descriptions
                        - Filtering out irrelevant administrative information
                        - Understanding legal document structure and terminology
                        - Handling multiple dates in the same context correctly
                        - Processing multiple pages efficiently in batch
                        
                        Always focus on the substance of legal proceedings rather than document formatting or metadata."""
                    },
                    {
                        "role": "user",
                        "content": batch_prompt
                    }
                ],
                model="deepseek-r1-distill-llama-70b",
                temperature=0.01,
                max_tokens=4000
            )
            
            all_entries = []
            content = response.choices[0].message.content
            self.logger.info(f"Received batch response")
            
            # Parse batch response
            current_page = None
            for line in content.split('\n'):
                line = line.strip()
                
                # Check for page markers
                if line.startswith('PAGE ') and ':' in line:
                    try:
                        current_page = int(line.split('PAGE ')[1].split(':')[0].strip())
                    except:
                        continue
                
                # Process date entries
                if '|' in line and current_page is not None:
                    date_str, event = line.split('|', 1)
                    date_str = date_str.strip()
                    event = event.strip()
                    
                    # Clean the date string
                    cleaned_date = self.clean_date_string(date_str)
                    
                    if not cleaned_date:
                        continue
                    
                    # Filter out metadata dates
                    event_lower = event.lower()
                    if any(keyword in event_lower for keyword in ['digitally signed', 'signing date', 'generated on', 'printed on', 'timestamp']):
                        continue
                    
                    try:
                        # Check if it's a year-only format
                        if re.match(r'^\d{4}$', cleaned_date):
                            all_entries.append({
                                'date': cleaned_date,
                                'event': event,
                                'page': current_page
                            })
                        else:
                            # Try to parse as complete date
                            parsed_date = None
                            date_formats = [
                                '%d/%m/%Y',
                                '%d-%m-%Y',
                                '%Y-%m-%d',
                                '%d %B %Y',
                                '%d %b %Y',
                                '%B %d, %Y',
                                '%b %d, %Y'
                            ]
                            
                            clean_date_str = re.sub(r'(\d+)(th|st|nd|rd)\s+', r'\1 ', cleaned_date)
                            clean_date_str = clean_date_str.strip()
                            
                            for date_format in date_formats:
                                try:
                                    parsed_date = datetime.strptime(clean_date_str, date_format)
                                    break
                                except ValueError:
                                    continue
                            
                            if parsed_date:
                                formatted_date = parsed_date.strftime('%d/%m/%Y')
                                all_entries.append({
                                    'date': formatted_date,
                                    'event': event,
                                    'page': current_page
                                })
                            else:
                                if re.match(r'^\d{4}$', clean_date_str):
                                    all_entries.append({
                                        'date': clean_date_str,
                                        'event': event,
                                        'page': current_page
                                    })
                    
                    except Exception as e:
                        self.logger.warning(f"Failed to parse date '{date_str}': {str(e)}")
                        continue
            
            self.logger.info(f"Extracted {len(all_entries)} entries from batch of {len(page_data_list)} pages")
            return all_entries
            
        except Exception as e:
            self.logger.error(f"Error processing batch: {str(e)}")
            return []

    def extract_dates_and_events(self, text: str, page_number: int) -> List[Dict]:
        """Legacy method for single page processing - kept for compatibility"""
        return self.extract_dates_and_events_batch([{'text': text, 'page_number': page_number}])

    def extract_text_with_ocr(self, pdf_path: str, page_number: int) -> str:
        """Extract text from PDF page using OCR when normal extraction fails"""
        try:
            self.logger.info(f"Attempting text extraction for page {page_number}")
            doc = fitz.open(pdf_path)
            page = doc[page_number - 1]
            
            # First try normal text extraction
            text = page.get_text()
            
            # If text extraction returns very little text, use OCR
            if len(text.strip()) < 50:
                self.logger.info(f"Normal text extraction failed, using OCR for page {page_number}")
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    temp_img_path = tmp_file.name
                    pix.save(temp_img_path)
                
                image = Image.open(temp_img_path)
                text = pytesseract.image_to_string(image)
                
                os.unlink(temp_img_path)
                self.logger.info(f"OCR extracted {len(text.strip())} characters from page {page_number}")
            else:
                self.logger.info(f"Normal text extraction succeeded for page {page_number}")
            
            doc.close()
            return text
            
        except Exception as e:
            self.logger.error(f"Error in text extraction for page {page_number}: {str(e)}")
            return ""

    def process_document(self, pdf_path: str) -> pd.DataFrame:
        try:
            self.logger.info(f"Starting to process document: {pdf_path}")
            
            doc = fitz.open(pdf_path)
            all_entries = []
            
            # Extract text from all pages first
            page_data = []
            for page_num in range(1, doc.page_count + 1):
                self.logger.info(f"Extracting text from page {page_num}/{doc.page_count}")
                text = self.extract_text_with_ocr(pdf_path, page_num)
                
                if text.strip():
                    page_data.append({
                        'page_number': page_num,
                        'text': text
                    })
                else:
                    self.logger.warning(f"Page {page_num} appears to be empty or unreadable")
            
            doc.close()
            
            # Process pages in batches of 3
            batch_size = 3
            for i in range(0, len(page_data), batch_size):
                batch = page_data[i:i + batch_size]
                self.logger.info(f"Processing batch {i//batch_size + 1} with {len(batch)} pages")
                
                entries = self.extract_dates_and_events_batch(batch)
                all_entries.extend(entries)
                
                self.logger.info(f"Found {len(entries)} entries in batch {i//batch_size + 1}")
            
            if not all_entries:
                self.logger.warning("No substantive date entries found in document")
                return pd.DataFrame(columns=['Date', 'Event Details', 'Source Pages', 'Page Numbers'])
            
            # Combine entries by date with improved handling
            self.logger.info("Combining entries by date")
            events_by_date = {}
            for entry in all_entries:
                date = entry['date']
                if date not in events_by_date:
                    events_by_date[date] = {'events': set(), 'pages': set()}
                events_by_date[date]['events'].add(entry['event'])
                events_by_date[date]['pages'].add(entry['page'])

            # Create DataFrame with improved formatting
            df_data = [
                {
                    'Date': date,
                    'Event Details': '\n'.join(sorted(info['events'])),
                    'Source Pages': f"Pages {', '.join(map(str, sorted(info['pages'])))}",
                    'Page Numbers': sorted(info['pages'])
                }
                for date, info in events_by_date.items()
            ]

            df = pd.DataFrame(df_data)
            
            # Enhanced sorting that handles both years and complete dates
            try:
                def sort_key(date_str):
                    # If it's just a year, convert to datetime with Jan 1
                    if re.match(r'^\d{4}$', date_str):
                        return datetime(int(date_str), 1, 1)
                    # Try to parse as complete date
                    try:
                        return pd.to_datetime(date_str, format='%d/%m/%Y')
                    except:
                        return datetime(1900, 1, 1)  # Default for unparseable dates
                
                df['sort_key'] = df['Date'].apply(sort_key)
                df = df.sort_values('sort_key')
                df = df.drop('sort_key', axis=1)
                
            except Exception as e:
                self.logger.warning(f"Could not sort dates: {str(e)}")

            self.logger.info(f"Successfully created timeline with {len(df)} substantive entries")
            return df

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            return pd.DataFrame(columns=['Date', 'Event Details', 'Source Pages', 'Page Numbers'])

def format_timeline_for_api(df: pd.DataFrame) -> List[Dict[str, Any]]:
    """Format timeline DataFrame for API response"""
    timeline_data = []
    
    for _, row in df.iterrows():
        # Extract page numbers from the string or use the direct list if available
        page_numbers = row.get('Page Numbers', [])
        if isinstance(page_numbers, str):
            # Try to parse page numbers from string format like "Pages 1, 2, 3"
            try:
                page_str = page_numbers.replace("Pages ", "")
                page_numbers = [int(p.strip()) for p in page_str.split(",")]
            except:
                page_numbers = []
        
        # Ensure page_numbers is a list
        if not isinstance(page_numbers, list):
            page_numbers = [page_numbers] if page_numbers else []
        
        entry = {
            "date": str(row['Date']),
            "event_details": str(row['Event Details']),
            "source_document": str(row.get('Source Document', 'Unknown')),
            "page_numbers": page_numbers
        }
        timeline_data.append(entry)
    
    return timeline_data

def create_word_document(df: pd.DataFrame, title: str = "Timeline of Events") -> Document:
    """Create a formatted Word document with the timeline"""
    doc = Document()

    # Add title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Add spacing

    # Check if DataFrame is empty
    if df.empty:
        # Add message for empty timeline
        empty_para = doc.add_paragraph()
        empty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        empty_run = empty_para.add_run("No timeline events were detected in this document.")
        empty_run.font.italic = True
        empty_run.font.size = Pt(12)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("End of Timeline")
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        
        return doc

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Set headers
    headers = ['Date', 'Event Details', 'Source Document', 'Source Pages']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)

    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Date'])
        row_cells[1].text = str(row['Event Details']).replace('•', '-')
        
        # Handle Source Document column
        if 'Source Document' in row:
            row_cells[2].text = str(row['Source Document'])
        else:
            row_cells[2].text = "Unknown"
            
        row_cells[3].text = str(row['Source Pages'])
        
        # Style data cells
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)  # Date
        row.cells[1].width = Inches(3.5)  # Event Details
        row.cells[2].width = Inches(1.5)  # Source Document
        row.cells[3].width = Inches(1.0)  # Source Pages

    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("End of Timeline")
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True

    return doc

def get_pdf_preview(pdf_path: str, page_number: int) -> str:
    """Generate PDF page preview and return as base64 string"""
    doc = None
    temp_img_path = None
    try:
        # Open the PDF
        doc = fitz.open(pdf_path)
        
        # Ensure valid page number
        if page_number < 1 or page_number > doc.page_count:
            page_number = 1
        
        # Get the specific page
        page = doc[page_number - 1]
        
        # Set a higher zoom factor for better quality
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        
        # Get the page pixmap
        pix = page.get_pixmap(matrix=mat)
        
        # Create a temporary file for the image
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            temp_img_path = tmp_file.name
            # Save the pixmap to the temporary file
            pix.save(temp_img_path)
        
        # Read the saved image and convert to base64
        with open(temp_img_path, 'rb') as img_file:
            base64_data = base64.b64encode(img_file.read()).decode()
        
        return base64_data
        
    except Exception as e:
        logger.error(f"Error generating PDF preview for page {page_number}: {str(e)}")
        return None
        
    finally:
        # Clean up resources
        if doc:
            doc.close()
        if temp_img_path and os.path.exists(temp_img_path):
            try:
                os.unlink(temp_img_path)
            except Exception as e:
                logger.error(f"Error cleaning up temporary file: {str(e)}")

def process_timeline(pdf_path: str, document_name: str = None) -> tuple:
    """Process a single PDF document and return timeline data and document"""
    try:
        extractor = TimelineExtractor()
        df = extractor.process_document(pdf_path)
        
        if document_name and not df.empty:
            df['Source Document'] = document_name
        
        doc = create_word_document(df, f"Timeline - {document_name if document_name else 'Document'}")
        
        return df, doc
    except Exception as e:
        logger.error(f"Error processing timeline: {str(e)}")
        return None, None

def validate_pdf_path(pdf_path: str) -> bool:
    """Validate if the PDF file exists and is accessible"""
    try:
        doc = fitz.open(pdf_path)
        doc.close()
        return True
    except Exception as e:
        logger.error(f"Error validating PDF path: {str(e)}")
        return False

def get_pdf_page_count(pdf_path: str) -> int:
    """Get the total number of pages in a PDF"""
    try:
        doc = fitz.open(pdf_path)
        count = doc.page_count
        doc.close()
        return count
    except Exception as e:
        logger.error(f"Error getting PDF page count: {str(e)}")
        return 0

def get_pdf_files_from_folder(folder_path: str) -> List[str]:
    """Get all PDF files from a folder"""
    try:
        pdf_files = []
        folder_path = Path(folder_path)
        
        if not folder_path.exists():
            logger.error(f"Folder does not exist: {folder_path}")
            return []
        
        # Find all PDF files in the folder
        for pdf_file in folder_path.glob("*.pdf"):
            pdf_files.append(str(pdf_file))
        
        # Also check for PDF files in subfolders
        for pdf_file in folder_path.glob("**/*.pdf"):
            if str(pdf_file) not in pdf_files:
                pdf_files.append(str(pdf_file))
        
        logger.info(f"Found {len(pdf_files)} PDF files in folder: {folder_path}")
        return sorted(pdf_files)
    
    except Exception as e:
        logger.error(f"Error getting PDF files from folder: {str(e)}")
        return []

def get_file_info(pdf_path: str) -> Dict[str, Any]:
    """Get basic information about a PDF file"""
    try:
        file_stat = os.stat(pdf_path)
        file_size = file_stat.st_size / (1024 * 1024)  # Size in MB
        modified_time = datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d %H:%M')
        
        # Get page count
        page_count = get_pdf_page_count(pdf_path)
        
        return {
            'name': os.path.basename(pdf_path),
            'size_mb': round(file_size, 2),
            'modified': modified_time,
            'pages': page_count,
            'path': pdf_path
        }
    except Exception as e:
        logger.error(f"Error getting file info for {pdf_path}: {str(e)}")
        return {
            'name': os.path.basename(pdf_path),
            'size_mb': 0,
            'modified': 'Unknown',
            'pages': 0,
            'path': pdf_path
        }

def display_file_selection_menu(pdf_files: List[str]) -> List[str]:
    """Display file selection menu and return selected files"""
    print("\n" + "=" * 80)
    print("                    📁 PDF FILE SELECTION MENU")
    print("=" * 80)
    
    # Get file information for all PDFs
    file_info_list = []
    for pdf_file in pdf_files:
        info = get_file_info(pdf_file)
        file_info_list.append(info)
    
    # Display files with detailed information
    print(f"{'No.':<4} {'File Name':<40} {'Size (MB)':<10} {'Pages':<8} {'Modified':<20}")
    print("-" * 80)
    
    for i, info in enumerate(file_info_list, 1):
        print(f"{i:<4} {info['name'][:38]:<40} {info['size_mb']:<10} {info['pages']:<8} {info['modified']:<20}")
    
    print("-" * 80)
    print(f"Total: {len(pdf_files)} PDF files found")
    print("\n📋 Selection Options:")
    print("1. Enter specific file numbers (e.g., 1,3,5 or 1-5)")
    print("2. Type 'all' to select all files")
    print("3. Type 'quit' to cancel")
    
    while True:
        try:
            selection = input("\n🔍 Enter your selection: ").strip().lower()
            
            if selection == 'quit':
                return []
            
            if selection == 'all':
                return pdf_files
            
            # Parse number selections
            selected_indices = set()
            
            # Split by comma and process each part
            parts = selection.split(',')
            for part in parts:
                part = part.strip()
                if '-' in part:
                    # Handle range (e.g., 1-5)
                    try:
                        start, end = map(int, part.split('-'))
                        if start < 1 or end > len(pdf_files) or start > end:
                            raise ValueError("Invalid range")
                        selected_indices.update(range(start, end + 1))
                    except ValueError:
                        print(f"❌ Invalid range format: {part}")
                        continue
                else:
                    # Handle single number
                    try:
                        num = int(part)
                        if num < 1 or num > len(pdf_files):
                            raise ValueError("Number out of range")
                        selected_indices.add(num)
                    except ValueError:
                        print(f"❌ Invalid number: {part}")
                        continue
            
            if not selected_indices:
                print("❌ No valid selections made. Please try again.")
                continue
            
            # Convert indices to file paths
            selected_files = [pdf_files[i-1] for i in sorted(selected_indices)]
            
            # Confirm selection
            print(f"\n✅ Selected {len(selected_files)} files:")
            for i, file_path in enumerate(selected_files, 1):
                print(f"  {i}. {os.path.basename(file_path)}")
            
            confirm = input(f"\n❓ Process these {len(selected_files)} files? (y/n): ").strip().lower()
            if confirm == 'y':
                return selected_files
            else:
                print("Selection cancelled. Please choose again.")
                continue
                
        except KeyboardInterrupt:
            print("\n\n❌ Operation cancelled by user.")
            return []
        except Exception as e:
            print(f"❌ Error in selection: {str(e)}. Please try again.")
            continue

def process_multiple_pdfs(pdf_paths: List[str], output_folder: str = "output") -> Dict[str, Any]:
    """Process multiple PDF files and generate combined timeline"""
    try:
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        all_dataframes = []
        processed_files = []
        failed_files = []
        
        logger.info(f"Processing {len(pdf_paths)} PDF files...")
        
        for i, pdf_path in enumerate(pdf_paths, 1):
            try:
                logger.info(f"Processing file {i}/{len(pdf_paths)}: {os.path.basename(pdf_path)}")
                
                # Validate PDF
                if not validate_pdf_path(pdf_path):
                    logger.error(f"Invalid PDF file: {pdf_path}")
                    failed_files.append(pdf_path)
                    continue
                
                # Process the PDF
                document_name = os.path.basename(pdf_path)
                df, doc = process_timeline(pdf_path, document_name)
                
                if df is not None and not df.empty:
                    all_dataframes.append(df)
                    processed_files.append(pdf_path)
                    
                    # Save individual Word document
                    individual_output_path = os.path.join(output_folder, f"timeline_{document_name.replace('.pdf', '.docx')}")
                    doc.save(individual_output_path)
                    logger.info(f"Saved individual timeline: {individual_output_path}")
                else:
                    logger.warning(f"No timeline data extracted from: {pdf_path}")
                    failed_files.append(pdf_path)
                    
            except Exception as e:
                logger.error(f"Error processing {pdf_path}: {str(e)}")
                failed_files.append(pdf_path)
        
        # Combine all dataframes
        if all_dataframes:
            combined_df = pd.concat(all_dataframes, ignore_index=True)
            
            # Sort by date using the enhanced sorting function
            try:
                def sort_key(date_str):
                    # If it's just a year, convert to datetime with Jan 1
                    if re.match(r'^\d{4}$', date_str):
                        return datetime(int(date_str), 1, 1)
                    # Try to parse as complete date
                    try:
                        return pd.to_datetime(date_str, format='%d/%m/%Y')
                    except:
                        return datetime(1900, 1, 1)  # Default for unparseable dates
                
                combined_df['sort_key'] = combined_df['Date'].apply(sort_key)
                combined_df = combined_df.sort_values('sort_key')
                combined_df = combined_df.drop('sort_key', axis=1)
                
            except Exception as e:
                logger.warning(f"Could not sort combined dates: {str(e)}")
            
            # Create combined Word document
            combined_doc = create_word_document(combined_df, "Combined Timeline - All Documents")
            combined_output_path = os.path.join(output_folder, "combined_timeline.docx")
            combined_doc.save(combined_output_path)
            
            # Save combined CSV
            csv_output_path = os.path.join(output_folder, "combined_timeline.csv")
            combined_df.to_csv(csv_output_path, index=False)
            
            logger.info(f"Saved combined timeline: {combined_output_path}")
            logger.info(f"Saved combined CSV: {csv_output_path}")
            
            return {
                "success": True,
                "processed_files": len(processed_files),
                "failed_files": len(failed_files),
                "total_events": len(combined_df),
                "output_folder": output_folder,
                "combined_timeline_path": combined_output_path,
                "csv_path": csv_output_path,
                "failed_file_list": failed_files
            }
        else:
            logger.error("No timeline data was extracted from any PDF files")
            return {
                "success": False,
                "error": "No timeline data extracted from any files",
                "failed_files": len(failed_files),
                "failed_file_list": failed_files
            }
            
    except Exception as e:
        logger.error(f"Error in process_multiple_pdfs: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

def main():
    """Main function with enhanced user interface for PDF timeline extraction"""
    print("=" * 60)
    print("           ENHANCED PDF TIMELINE EXTRACTOR")
    print("=" * 60)
    print("This tool extracts substantive chronological events from PDF documents")
    print("while filtering out metadata and page corner information.\n")
    
    while True:
        print("Choose an option:")
        print("1. Process a single PDF file")
        print("2. Browse and select PDFs from a folder")
        print("3. Process all PDFs in a folder")
        print("4. Exit")
        
        choice = input("\nEnter your choice (1-4): ").strip()
        
        if choice == "1":
            # Single PDF processing
            pdf_path = input("\nEnter the path to your PDF file: ").strip().strip('"')
            
            if not os.path.exists(pdf_path):
                print(f"❌ Error: File not found - {pdf_path}")
                continue
            
            if not pdf_path.lower().endswith('.pdf'):
                print("❌ Error: Please provide a PDF file")
                continue
            
            print(f"\n🔄 Processing PDF: {os.path.basename(pdf_path)}")
            
            try:
                # Validate PDF
                if not validate_pdf_path(pdf_path):
                    print("❌ Error: Invalid or corrupted PDF file")
                    continue
                
                # Get page count
                page_count = get_pdf_page_count(pdf_path)
                print(f"📄 PDF has {page_count} pages")
                
                # Process the PDF
                document_name = os.path.basename(pdf_path)
                df, doc = process_timeline(pdf_path, document_name)
                
                if df is not None and not df.empty:
                    # Create output folder
                    output_folder = "timeline_output"
                    os.makedirs(output_folder, exist_ok=True)
                    
                    # Save Word document
                    output_path = os.path.join(output_folder, f"timeline_{document_name.replace('.pdf', '.docx')}")
                    doc.save(output_path)
                    
                    # Save CSV
                    csv_path = os.path.join(output_folder, f"timeline_{document_name.replace('.pdf', '.csv')}")
                    df.to_csv(csv_path, index=False)
                    
                    print(f"✅ Success! Extracted {len(df)} substantive timeline events")
                    print(f"📁 Word document saved: {output_path}")
                    print(f"📊 CSV file saved: {csv_path}")
                    
                    # Show preview of results
                    print(f"\n📋 Preview of extracted events:")
                    print("-" * 80)
                    for i, (_, row) in enumerate(df.head(3).iterrows()):
                        print(f"{i+1}. {row['Date']}: {row['Event Details'][:150]}...")
                    
                    if len(df) > 3:
                        print(f"... and {len(df) - 3} more events")
                
                else:
                    print("❌ No substantive timeline events were found in this PDF")
                    print("This could be due to:")
                    print("- No case-related dates present in the document")
                    print("- Document contains only metadata/administrative information")
                    print("- Scanned document with poor OCR quality")
                
            except Exception as e:
                print(f"❌ Error processing PDF: {str(e)}")
        
        elif choice == "2":
            # Browse and select specific PDFs from folder
            folder_path = input("\nEnter the path to your folder containing PDFs: ").strip().strip('"')
            
            if not os.path.exists(folder_path):
                print(f"❌ Error: Folder not found - {folder_path}")
                continue
            
            if not os.path.isdir(folder_path):
                print("❌ Error: Please provide a valid folder path")
                continue
            
            print(f"\n🔍 Searching for PDF files in: {folder_path}")
            
            # Get all PDF files
            pdf_files = get_pdf_files_from_folder(folder_path)
            
            if not pdf_files:
                print("❌ No PDF files found in the specified folder")
                continue
            
            # Display selection menu and get user choice
            selected_files = display_file_selection_menu(pdf_files)
            
            if not selected_files:
                print("❌ No files selected or operation cancelled")
                continue
            
            print(f"\n🔄 Processing {len(selected_files)} selected PDF files...")
            
            # Process selected PDFs
            output_folder = f"timeline_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            result = process_multiple_pdfs(selected_files, output_folder)
            
            if result["success"]:
                print(f"✅ Processing completed successfully!")
                print(f"📊 Results:")
                print(f"  - Files processed: {result['processed_files']}")
                print(f"  - Files failed: {result['failed_files']}")
                print(f"  - Total events extracted: {result['total_events']}")
                print(f"📁 Output folder: {result['output_folder']}")
                print(f"📄 Combined timeline: {result['combined_timeline_path']}")
                print(f"📊 CSV export: {result['csv_path']}")
                
                if result['failed_file_list']:
                    print(f"\n⚠️  Failed files:")
                    for failed_file in result['failed_file_list']:
                        print(f"  - {os.path.basename(failed_file)}")
            else:
                print(f"❌ Processing failed: {result.get('error', 'Unknown error')}")
        
        elif choice == "3":
            # Process all PDFs in folder
            folder_path = input("\nEnter the path to your folder containing PDFs: ").strip().strip('"')
            
            if not os.path.exists(folder_path):
                print(f"❌ Error: Folder not found - {folder_path}")
                continue
            
            if not os.path.isdir(folder_path):
                print("❌ Error: Please provide a valid folder path")
                continue
            
            print(f"\n🔍 Searching for PDF files in: {folder_path}")
            
            # Get all PDF files
            pdf_files = get_pdf_files_from_folder(folder_path)
            
            if not pdf_files:
                print("❌ No PDF files found in the specified folder")
                continue
            
            print(f"📚 Found {len(pdf_files)} PDF files:")
            for i, pdf_file in enumerate(pdf_files, 1):
                print(f"  {i}. {os.path.basename(pdf_file)}")
            
            confirm = input(f"\nProcess all {len(pdf_files)} PDF files? (y/n): ").strip().lower()
            if confirm != 'y':
                print("Operation cancelled")
                continue
            
            print(f"\n🔄 Processing {len(pdf_files)} PDF files...")
            
            # Process all PDFs
            output_folder = f"timeline_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            result = process_multiple_pdfs(pdf_files, output_folder)
            
            if result["success"]:
                print(f"✅ Processing completed successfully!")
                print(f"📊 Results:")
                print(f"  - Files processed: {result['processed_files']}")
                print(f"  - Files failed: {result['failed_files']}")
                print(f"  - Total events extracted: {result['total_events']}")
                print(f"📁 Output folder: {result['output_folder']}")
                print(f"📄 Combined timeline: {result['combined_timeline_path']}")
                print(f"📊 CSV export: {result['csv_path']}")
                
                if result['failed_file_list']:
                    print(f"\n⚠️  Failed files:")
                    for failed_file in result['failed_file_list']:
                        print(f"  - {os.path.basename(failed_file)}")
            else:
                print(f"❌ Processing failed: {result.get('error', 'Unknown error')}")
        
        elif choice == "4":
            print("\n👋 Thank you for using Enhanced PDF Timeline Extractor!")
            break
        
        else:
            print("❌ Invalid choice. Please enter 1, 2, 3, or 4.")
        
        print("\n" + "=" * 60)

if __name__ == "__main__":
    main()

